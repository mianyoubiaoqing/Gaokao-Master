"""Local knowledge base pipeline for Gaokao-Master.

The manager keeps a plain-text-first, Obsidian-like workspace while also
maintaining a ChromaDB semantic index for retrieval.
"""

from __future__ import annotations

import hashlib
import math
import re
import shutil
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from uuid import uuid4

import chromadb
import fitz
from chromadb.utils import embedding_functions
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from loguru import logger


SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}


class DocumentExtractionError(ValueError):
    """Raised when a document cannot produce useful indexable text."""

    def __init__(self, message: str, diagnostic_path: Path | None = None) -> None:
        super().__init__(message)
        self.diagnostic_path = diagnostic_path


class LocalHashEmbeddingFunction:
    """Small offline embedding function that never downloads model files.

    It is not as semantically rich as a transformer model, but it is stable,
    fast, Chinese-aware, and good enough as a zero-setup fallback. Keyword BM25
    retrieval still complements it in `fuzzy_retrieve`.
    """

    def __init__(self, dimensions: int = 384) -> None:
        self.dimensions = dimensions

    def __call__(self, input: list[str]) -> list[list[float]]:  # noqa: A002
        return [self._embed_text(text) for text in input]

    @staticmethod
    def name() -> str:
        return "gaokao-master-local-hash"

    @staticmethod
    def build_from_config(config: dict[str, Any]) -> "LocalHashEmbeddingFunction":
        return LocalHashEmbeddingFunction(
            dimensions=int(config.get("dimensions", 384))
        )

    def get_config(self) -> dict[str, Any]:
        return {"dimensions": self.dimensions}

    def default_space(self) -> str:
        return "cosine"

    def supported_spaces(self) -> list[str]:
        return ["cosine", "l2", "ip"]

    def _embed_text(self, text: str) -> list[float]:
        vector = [0.0] * self.dimensions
        tokens = re.findall(r"[\u4e00-\u9fff]|[A-Za-z0-9_]+", text.lower())
        if not tokens:
            return vector

        for token in tokens:
            digest = hashlib.blake2b(token.encode("utf-8"), digest_size=8).digest()
            bucket = int.from_bytes(digest[:4], "little") % self.dimensions
            sign = 1.0 if digest[4] % 2 == 0 else -1.0
            vector[bucket] += sign

        norm = math.sqrt(sum(value * value for value in vector))
        if norm == 0:
            return vector
        return [value / norm for value in vector]


@dataclass(frozen=True)
class KBPaths:
    """Resolved filesystem paths used by the knowledge base."""

    root: Path
    raw: Path
    vector_store: Path


@dataclass(frozen=True)
class IngestedDocument:
    """Result returned after a source document is ingested."""

    source_path: Path
    markdown_path: Path
    chunk_count: int
    collection_name: str


class KnowledgeBaseManager:
    """Manage local Gaokao knowledge files and their vector index.

    Directory layout:

    ```text
    Gaokao_KB/
      _raw/
      .vector_store/
      Chinese/
        古诗文/
          example.md
    ```

    Subject and topic folders are intentionally human-readable so that the KB
    remains easy to browse and edit with tools such as Obsidian or VS Code.
    """

    def __init__(
        self,
        kb_root: str | Path = "Gaokao_KB",
        collection_name: str | None = None,
        embedding_model_name: str = "local-hash",
        chunk_size: int = 1_200,
        chunk_overlap: int = 120,
        ocr_client: Any | None = None,
        ocr_dpi: int = 180,
        ocr_max_pages: int = 20,
        ocr_extract_media: bool = True,
    ) -> None:
        self.paths = self._init_paths(Path(kb_root))
        self.collection_name = collection_name or self._default_collection_name(
            embedding_model_name
        )
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.ocr_client = ocr_client
        self.ocr_dpi = ocr_dpi
        self.ocr_max_pages = ocr_max_pages
        self.ocr_extract_media = ocr_extract_media

        embedding_fn = self._build_embedding_function(embedding_model_name)
        self.client = chromadb.PersistentClient(path=str(self.paths.vector_store))
        self.collection = self.client.get_or_create_collection(
            name=self.collection_name,
            embedding_function=embedding_fn,
            metadata={"description": "Gaokao-Master local knowledge base"},
        )

    def ingest_file(
        self,
        file_path: str | Path,
        subject: str,
        topic: str,
        *,
        title: str | None = None,
        tags: list[str] | None = None,
        source_url: str | None = None,
        keep_raw_copy: bool = True,
    ) -> IngestedDocument:
        """Parse a supported file, save Markdown, then upsert chunks to ChromaDB."""

        source_path = Path(file_path).expanduser().resolve()
        if not source_path.exists():
            raise FileNotFoundError(f"Source file does not exist: {source_path}")

        extension = source_path.suffix.lower()
        if extension not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type {extension!r}; expected one of "
                f"{sorted(SUPPORTED_EXTENSIONS)}."
            )

        if keep_raw_copy and extension in {".pdf", ".docx"}:
            self._copy_raw_file(source_path, subject, topic)

        markdown_path = self._convert_to_markdown(
            source_path=source_path,
            subject=subject,
            topic=topic,
            title=title,
            tags=tags or [],
            source_url=source_url,
        )
        chunks = self.chunk_markdown(markdown_path.read_text(encoding="utf-8"))
        self.upsert_markdown(
            markdown_path=markdown_path,
            chunks=chunks,
            subject=subject,
            topic=topic,
            source_path=source_path,
            tags=tags or [],
            source_url=source_url,
        )

        return IngestedDocument(
            source_path=source_path,
            markdown_path=markdown_path,
            chunk_count=len(chunks),
            collection_name=self.collection_name,
        )

    def parse_pdf_to_md(
        self,
        pdf_path: str | Path,
        subject: str,
        topic: str,
        *,
        title: str | None = None,
        tags: list[str] | None = None,
        source_url: str | None = None,
    ) -> Path:
        """Extract PDF text with PyMuPDF and save it as Markdown.

        Scanned image-only PDFs cannot be reliably converted without OCR. Those
        files are rejected with a clear error so they do not pollute the vector
        index with empty chunks.
        """

        source_path = Path(pdf_path).expanduser().resolve()
        if source_path.suffix.lower() != ".pdf":
            raise ValueError(f"Expected a .pdf file, got: {source_path}")

        pages: list[str] = []
        empty_pages: list[int] = []
        with fitz.open(source_path) as pdf:
            for page_index, page in enumerate(pdf, start=1):
                text = self._extract_pdf_page_text(page)
                if text:
                    pages.append(f"## Page {page_index}\n\n{text}")
                else:
                    empty_pages.append(page_index)

        body = "\n\n".join(pages)
        meaningful_chars = len(re.sub(r"\s+", "", body))
        if meaningful_chars < 80:
            if self._can_use_ocr():
                logger.info("PDF text is sparse; starting OCR for {}", source_path)
                body = self._ocr_pdf_to_markdown(source_path)
                meaningful_chars = len(re.sub(r"\s+", "", body))
                if meaningful_chars >= 80:
                    return self._write_markdown(
                        body=body,
                        source_path=source_path,
                        subject=subject,
                        topic=topic,
                        title=title,
                        tags=[*(tags or []), "ocr"],
                        source_url=source_url,
                    )

            diagnostic = self._build_empty_pdf_diagnostic(source_path, empty_pages)
            diagnostic_path = self._write_markdown(
                body=diagnostic,
                source_path=source_path,
                subject=subject,
                topic=topic,
                title=f"{title or source_path.stem}_需OCR",
                tags=[*(tags or []), "needs-ocr"],
                source_url=source_url,
            )
            raise DocumentExtractionError(
                "PDF has too little extractable text. It is likely a scanned "
                "paper and needs OCR before indexing.",
                diagnostic_path=diagnostic_path,
            )

        return self._write_markdown(
            body=body,
            source_path=source_path,
            subject=subject,
            topic=topic,
            title=title,
            tags=tags or [],
            source_url=source_url,
        )

    def parse_docx_to_md(
        self,
        docx_path: str | Path,
        subject: str,
        topic: str,
        *,
        title: str | None = None,
        tags: list[str] | None = None,
        source_url: str | None = None,
    ) -> Path:
        """Extract DOCX paragraphs and tables, then save them as Markdown."""

        source_path = Path(docx_path).expanduser().resolve()
        if source_path.suffix.lower() != ".docx":
            raise ValueError(f"Expected a .docx file, got: {source_path}")

        document = Document(str(source_path))
        blocks: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = paragraph.style.name.lower() if paragraph.style else ""
            blocks.append(self._format_docx_paragraph(text, style_name))

        for table in document.tables:
            markdown_table = self._table_to_markdown(table)
            if markdown_table:
                blocks.append(markdown_table)

        body = "\n\n".join(blocks)
        return self._write_markdown(
            body=body,
            source_path=source_path,
            subject=subject,
            topic=topic,
            title=title,
            tags=tags or [],
            source_url=source_url,
        )

    def chunk_markdown(self, markdown_text: str) -> list[str]:
        """Chunk Markdown while trying to keep exam questions and answers intact.

        The method first splits the document into question-like blocks. If a
        block is too large, it falls back to a Markdown-aware recursive splitter.
        """

        normalized = self._normalize_text(markdown_text)
        question_blocks = self._split_question_blocks(normalized)

        if len(question_blocks) <= 1:
            return self._recursive_split(normalized)

        chunks: list[str] = []
        buffer: list[str] = []
        buffer_size = 0

        for block in question_blocks:
            block = block.strip()
            if not block:
                continue

            if len(block) > self.chunk_size:
                if buffer:
                    chunks.append("\n\n".join(buffer).strip())
                    buffer = []
                    buffer_size = 0
                chunks.extend(self._recursive_split(block))
                continue

            next_size = buffer_size + len(block)
            if buffer and next_size > self.chunk_size:
                chunks.append("\n\n".join(buffer).strip())
                buffer = [block]
                buffer_size = len(block)
            else:
                buffer.append(block)
                buffer_size = next_size

        if buffer:
            chunks.append("\n\n".join(buffer).strip())

        return [chunk for chunk in chunks if chunk]

    def upsert_markdown(
        self,
        markdown_path: str | Path,
        chunks: list[str],
        subject: str,
        topic: str,
        *,
        source_path: str | Path | None = None,
        tags: list[str] | None = None,
        source_url: str | None = None,
    ) -> None:
        """Upsert Markdown chunks into the local ChromaDB collection."""

        md_path = Path(markdown_path).resolve()
        if not chunks:
            logger.warning("No chunks generated for {}", md_path)
            return

        document_id = self._stable_id(str(md_path))
        ids: list[str] = []
        metadatas: list[dict[str, Any]] = []

        for index, chunk in enumerate(chunks):
            ids.append(f"{document_id}:{index:04d}")
            metadatas.append(
                {
                    "document_id": document_id,
                    "chunk_index": index,
                    "subject": subject,
                    "topic": topic,
                    "markdown_path": str(md_path),
                    "source_path": str(source_path or md_path),
                    "source_url": source_url or "",
                    "tags": ",".join(tags or []),
                    "updated_at": datetime.now(timezone.utc).isoformat(),
                }
            )

        self.collection.upsert(ids=ids, documents=chunks, metadatas=metadatas)
        logger.info("Upserted {} chunks for {}", len(chunks), md_path)

    def rebuild_index(self) -> int:
        """Rebuild the vector index from all Markdown files under the KB root."""

        markdown_files = [
            path
            for path in self.paths.root.rglob("*.md")
            if not self._is_hidden_or_internal(path)
        ]

        total_chunks = 0
        for markdown_path in markdown_files:
            relative = markdown_path.relative_to(self.paths.root)
            if len(relative.parts) < 3:
                logger.warning("Skipping Markdown outside subject/topic: {}", markdown_path)
                continue

            subject, topic = relative.parts[0], relative.parts[1]
            text = markdown_path.read_text(encoding="utf-8")
            chunks = self.chunk_markdown(text)
            self.upsert_markdown(
                markdown_path=markdown_path,
                chunks=chunks,
                subject=subject,
                topic=topic,
            )
            total_chunks += len(chunks)

        return total_chunks

    def delete_vectors_for_markdown(self, markdown_path: str | Path) -> int:
        """Delete vector records generated from one Markdown file."""

        md_path = Path(markdown_path).resolve()
        return self._delete_vectors_where({"markdown_path": str(md_path)})

    def delete_vectors_for_source(self, source_path: str | Path) -> int:
        """Delete vector records generated from one raw source file."""

        src_path = Path(source_path).resolve()
        return self._delete_vectors_where({"source_path": str(src_path)})

    def clear_vector_collection(self) -> str:
        """Delete the current Chroma collection."""

        self.client.delete_collection(self.collection_name)
        return self.collection_name

    def _convert_to_markdown(
        self,
        source_path: Path,
        subject: str,
        topic: str,
        title: str | None,
        tags: list[str],
        source_url: str | None,
    ) -> Path:
        extension = source_path.suffix.lower()
        if extension == ".pdf":
            return self.parse_pdf_to_md(
                source_path,
                subject,
                topic,
                title=title,
                tags=tags,
                source_url=source_url,
            )
        if extension == ".docx":
            return self.parse_docx_to_md(
                source_path,
                subject,
                topic,
                title=title,
                tags=tags,
                source_url=source_url,
            )

        body = source_path.read_text(encoding="utf-8", errors="ignore")
        if extension == ".txt":
            body = f"```text\n{body.strip()}\n```"

        return self._write_markdown(
            body=body,
            source_path=source_path,
            subject=subject,
            topic=topic,
            title=title,
            tags=tags,
            source_url=source_url,
        )

    @staticmethod
    def _init_paths(root: Path) -> KBPaths:
        resolved_root = root.expanduser().resolve()
        raw = resolved_root / "_raw"
        vector_store = resolved_root / ".vector_store"

        for path in (resolved_root, raw, vector_store):
            path.mkdir(parents=True, exist_ok=True)

        return KBPaths(root=resolved_root, raw=raw, vector_store=vector_store)

    def _copy_raw_file(self, source_path: Path, subject: str, topic: str) -> Path:
        target_dir = self.paths.raw / self._safe_part(subject) / self._safe_part(topic)
        target_dir.mkdir(parents=True, exist_ok=True)
        target = target_dir / source_path.name
        shutil.copy2(source_path, target)
        return target

    def _write_markdown(
        self,
        body: str,
        source_path: Path,
        subject: str,
        topic: str,
        title: str | None,
        tags: list[str],
        source_url: str | None,
    ) -> Path:
        title = title or source_path.stem
        subject_dir = self.paths.root / self._safe_part(subject) / self._safe_part(topic)
        subject_dir.mkdir(parents=True, exist_ok=True)

        markdown_path = subject_dir / f"{self._safe_filename(title)}.md"
        frontmatter = self._frontmatter(
            title=title,
            subject=subject,
            topic=topic,
            source_path=source_path,
            source_url=source_url,
            tags=tags,
        )
        markdown_path.write_text(
            f"{frontmatter}\n\n# {title}\n\n{body.strip()}\n",
            encoding="utf-8",
        )
        return markdown_path

    @staticmethod
    def _frontmatter(
        *,
        title: str,
        subject: str,
        topic: str,
        source_path: Path,
        source_url: str | None,
        tags: list[str],
    ) -> str:
        tag_values = ", ".join(f'"{tag}"' for tag in tags)
        source_url_line = f'source_url: "{source_url}"' if source_url else 'source_url: ""'
        return "\n".join(
            [
                "---",
                f'title: "{title}"',
                f'subject: "{subject}"',
                f'topic: "{topic}"',
                f'source_path: "{source_path}"',
                source_url_line,
                f"tags: [{tag_values}]",
                f'ingested_at: "{datetime.now(timezone.utc).isoformat()}"',
                "---",
            ]
        )

    @staticmethod
    def _format_docx_paragraph(text: str, style_name: str) -> str:
        if "heading 1" in style_name or "标题 1" in style_name:
            return f"# {text}"
        if "heading 2" in style_name or "标题 2" in style_name:
            return f"## {text}"
        if "heading 3" in style_name or "标题 3" in style_name:
            return f"### {text}"
        return text

    @staticmethod
    def _table_to_markdown(table: Any) -> str:
        rows = [
            [cell.text.strip().replace("\n", "<br>") for cell in row.cells]
            for row in table.rows
        ]
        rows = [row for row in rows if any(cell for cell in row)]
        if not rows:
            return ""

        header = rows[0]
        divider = ["---"] * len(header)
        body_rows = rows[1:] or [[""] * len(header)]

        def render_row(row: list[str]) -> str:
            padded = row + [""] * (len(header) - len(row))
            return "| " + " | ".join(padded[: len(header)]) + " |"

        return "\n".join(
            [render_row(header), render_row(divider)]
            + [render_row(row) for row in body_rows]
        )

    @staticmethod
    def _extract_pdf_page_text(page: fitz.Page) -> str:
        text = page.get_text("text", sort=True).strip()
        blocks = page.get_text("blocks", sort=True)
        block_texts = [
            str(block[4]).strip()
            for block in blocks
            if len(block) >= 5 and str(block[4]).strip()
        ]
        block_text = "\n\n".join(block_texts).strip()
        return block_text if len(block_text) > len(text) else text

    @staticmethod
    def _build_empty_pdf_diagnostic(source_path: Path, empty_pages: list[int]) -> str:
        page_hint = (
            "、".join(str(page) for page in empty_pages[:20])
            if empty_pages
            else "unknown"
        )
        return "\n".join(
            [
                "## PDF 解析诊断",
                "",
                "这份 PDF 没有检测到足够的可复制文本，因此没有写入向量索引。",
                "",
                f"- 原始文件：`{source_path}`",
                f"- 空文本页：{page_hint}",
                "",
                "常见原因：",
                "",
                "- 试卷是扫描图片版 PDF。",
                "- 下载到的文件不是实际试卷，可能是跳转页或防盗链页面。",
                "- PDF 中文本被转成图片或轮廓。",
                "",
                "处理建议：",
                "",
                "- 优先下载 DOCX 或可复制文字版 PDF。",
                "- 使用 OCR 工具先把扫描版 PDF 转成可复制文本 PDF。",
                "- 或把 OCR 后的 TXT/Markdown/DOCX 上传到资料导入页。",
            ]
        )

    def _can_use_ocr(self) -> bool:
        return bool(
            self.ocr_client
            and getattr(self.ocr_client, "is_configured", False)
        )

    def _ocr_pdf_to_markdown(self, source_path: Path) -> str:
        if not self._can_use_ocr():
            return ""

        page_markdown: list[str] = []
        with fitz.open(source_path) as pdf:
            total_pages = len(pdf)
            pages_to_process = min(total_pages, max(1, self.ocr_max_pages))
            matrix = fitz.Matrix(self.ocr_dpi / 72, self.ocr_dpi / 72)

            for page_index in range(pages_to_process):
                page = pdf[page_index]
                pixmap = page.get_pixmap(matrix=matrix, alpha=False)
                image_bytes = pixmap.tobytes("png")
                media_links = self._extract_pdf_page_media_links(
                    page=page,
                    source_path=source_path,
                    page_number=page_index + 1,
                    matrix=matrix,
                )
                try:
                    page_text = self.ocr_client.ocr_image(
                        image_bytes,
                        page_number=page_index + 1,
                        total_pages=total_pages,
                    )
                except Exception as exc:
                    logger.warning(
                        "OCR failed for {} page {}: {}",
                        source_path,
                        page_index + 1,
                        exc,
                    )
                    page_text = ""

                if page_text.strip():
                    page_text = self._insert_media_links_into_ocr_text(
                        page_text.strip(),
                        media_links,
                    )
                    page_markdown.append(
                        f"## Page {page_index + 1} OCR\n\n{page_text.strip()}"
                    )
                elif media_links:
                    page_markdown.append(
                        f"## Page {page_index + 1} OCR\n\n"
                        + "\n\n".join(media_links)
                    )

            if total_pages > pages_to_process:
                page_markdown.append(
                    "\n".join(
                        [
                            "## OCR 截断提示",
                            "",
                            f"该 PDF 共 {total_pages} 页，本次只 OCR 前 {pages_to_process} 页。",
                            "如需处理更多页面，请在 WebUI 中调高 OCR 最大页数。",
                        ]
                    )
                )

        return "\n\n".join(page_markdown).strip()

    def _extract_pdf_page_media_links(
        self,
        *,
        page: fitz.Page,
        source_path: Path,
        page_number: int,
        matrix: fitz.Matrix,
    ) -> list[str]:
        if not self.ocr_extract_media:
            return []

        rects = self._candidate_media_rects(page)
        if not rects:
            return []

        media_dir = (
            self.paths.root
            / "assets"
            / self._safe_filename(source_path.stem)
        )
        media_dir.mkdir(parents=True, exist_ok=True)

        links: list[str] = []
        for media_index, rect in enumerate(rects, start=1):
            target = media_dir / (
                f"page_{page_number:03d}_figure_{media_index:02d}.png"
            )
            target = self._unique_path(target)
            try:
                clip = self._expand_rect(rect, page.rect, padding=6)
                pixmap = page.get_pixmap(matrix=matrix, clip=clip, alpha=False)
                pixmap.save(str(target))
            except Exception as exc:
                logger.warning(
                    "Could not extract media from {} page {}: {}",
                    source_path,
                    page_number,
                    exc,
                )
                continue

            relative = target.relative_to(self.paths.root).as_posix()
            links.append(f"![page {page_number} figure {media_index}]({relative})")

        return links

    def _candidate_media_rects(self, page: fitz.Page) -> list[fitz.Rect]:
        page_rect = page.rect
        candidates: list[fitz.Rect] = []

        try:
            blocks = page.get_text("dict").get("blocks", [])
        except Exception:
            blocks = []

        for block in blocks:
            if block.get("type") != 1 or "bbox" not in block:
                continue
            rect = fitz.Rect(block["bbox"])
            if self._is_reasonable_media_rect(rect, page_rect):
                candidates.append(rect)

        drawing_rects: list[fitz.Rect] = []
        try:
            drawings = page.get_drawings()
        except Exception:
            drawings = []

        for drawing in drawings:
            rect = drawing.get("rect")
            if not rect:
                continue
            rect = fitz.Rect(rect)
            if rect.width >= 4 and rect.height >= 4:
                drawing_rects.append(self._expand_rect(rect, page_rect, padding=4))

        for rect in self._merge_nearby_rects(drawing_rects, padding=10):
            if self._is_reasonable_media_rect(rect, page_rect):
                candidates.append(rect)

        return self._dedupe_media_rects(candidates)

    @staticmethod
    def _is_reasonable_media_rect(rect: fitz.Rect, page_rect: fitz.Rect) -> bool:
        if rect.is_empty or rect.is_infinite:
            return False
        page_area = max(page_rect.get_area(), 1)
        area_ratio = rect.get_area() / page_area
        if area_ratio < 0.008 or area_ratio > 0.60:
            return False
        return rect.width >= 36 and rect.height >= 28

    @classmethod
    def _merge_nearby_rects(
        cls,
        rects: list[fitz.Rect],
        *,
        padding: float,
    ) -> list[fitz.Rect]:
        clusters: list[fitz.Rect] = []
        for rect in rects:
            expanded = fitz.Rect(
                rect.x0 - padding,
                rect.y0 - padding,
                rect.x1 + padding,
                rect.y1 + padding,
            )
            merged = False
            for index, cluster in enumerate(clusters):
                if cls._rects_overlap_or_touch(expanded, cluster):
                    clusters[index] = cluster | expanded
                    merged = True
                    break
            if not merged:
                clusters.append(expanded)

        changed = True
        while changed:
            changed = False
            merged_clusters: list[fitz.Rect] = []
            for rect in clusters:
                for index, existing in enumerate(merged_clusters):
                    if cls._rects_overlap_or_touch(rect, existing):
                        merged_clusters[index] = existing | rect
                        changed = True
                        break
                else:
                    merged_clusters.append(rect)
            clusters = merged_clusters
        return clusters

    @staticmethod
    def _rects_overlap_or_touch(a: fitz.Rect, b: fitz.Rect) -> bool:
        return not (
            a.x1 < b.x0
            or b.x1 < a.x0
            or a.y1 < b.y0
            or b.y1 < a.y0
        )

    @classmethod
    def _dedupe_media_rects(cls, rects: list[fitz.Rect]) -> list[fitz.Rect]:
        result: list[fitz.Rect] = []
        for rect in sorted(rects, key=lambda item: (item.y0, item.x0)):
            if any(cls._rect_overlap_ratio(rect, existing) > 0.72 for existing in result):
                continue
            result.append(rect)
        return result

    @staticmethod
    def _rect_overlap_ratio(a: fitz.Rect, b: fitz.Rect) -> float:
        intersection = fitz.Rect(
            max(a.x0, b.x0),
            max(a.y0, b.y0),
            min(a.x1, b.x1),
            min(a.y1, b.y1),
        )
        if intersection.is_empty:
            return 0.0
        return intersection.get_area() / max(min(a.get_area(), b.get_area()), 1)

    @staticmethod
    def _expand_rect(
        rect: fitz.Rect,
        page_rect: fitz.Rect,
        *,
        padding: float,
    ) -> fitz.Rect:
        expanded = fitz.Rect(
            rect.x0 - padding,
            rect.y0 - padding,
            rect.x1 + padding,
            rect.y1 + padding,
        )
        return expanded & page_rect

    @staticmethod
    def _insert_media_links_into_ocr_text(
        page_text: str,
        media_links: list[str],
    ) -> str:
        if not media_links:
            return page_text

        text = page_text
        marker = "[IMAGE_HERE]"
        used_count = 0
        for link in media_links:
            if marker not in text:
                break
            text = text.replace(marker, link, 1)
            used_count += 1

        remaining = media_links[used_count:]
        if remaining:
            text = text.rstrip() + "\n\n" + "\n\n".join(remaining)
        return text.replace(marker, "").strip()

    @staticmethod
    def _unique_path(path: Path) -> Path:
        if not path.exists():
            return path
        index = 2
        while True:
            candidate = path.with_name(f"{path.stem}_{index}{path.suffix}")
            if not candidate.exists():
                return candidate
            index += 1

    def _recursive_split(self, text: str) -> list[str]:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=["\n## ", "\n### ", "\n\n", "\n", "。", "；", "，", " ", ""],
        )
        return [chunk.strip() for chunk in splitter.split_text(text) if chunk.strip()]

    @staticmethod
    def _split_question_blocks(text: str) -> list[str]:
        question_pattern = re.compile(
            r"(?m)(?=^(?:"
            r"第\s*[一二三四五六七八九十百\d]+\s*[题问]|"
            r"[一二三四五六七八九十]+[、.．]|"
            r"\d{1,3}[、.．)]|"
            r"【(?:例题|真题|题目|答案|解析)】|"
            r"(?:答案|解析)[:：]"
            r"))"
        )
        parts = [part.strip() for part in question_pattern.split(text) if part.strip()]
        return parts or [text]

    @staticmethod
    def _normalize_text(text: str) -> str:
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[ \t]+", " ", text)
        return text.strip()

    @staticmethod
    def _safe_part(value: str) -> str:
        value = value.strip().replace("\\", "_").replace("/", "_")
        return re.sub(r'[<>:"|?*]+', "_", value) or "Uncategorized"

    @classmethod
    def _safe_filename(cls, value: str) -> str:
        filename = cls._safe_part(value)
        filename = re.sub(r"\s+", "_", filename)
        return filename[:120] or f"document_{uuid4().hex[:8]}"

    @staticmethod
    def _stable_id(value: str) -> str:
        return hashlib.sha256(value.encode("utf-8")).hexdigest()[:24]

    def _is_hidden_or_internal(self, path: Path) -> bool:
        relative_parts = path.relative_to(self.paths.root).parts
        return any(part.startswith(".") or part == "_raw" for part in relative_parts)

    def _delete_vectors_where(self, where: dict[str, str]) -> int:
        try:
            result = self.collection.get(where=where, include=["metadatas"])
        except Exception as exc:
            logger.warning("Could not list vectors for deletion {}: {}", where, exc)
            return 0

        ids = result.get("ids", [])
        if not ids:
            return 0

        self.collection.delete(ids=ids)
        logger.info("Deleted {} vector records for {}", len(ids), where)
        return len(ids)

    @staticmethod
    def _build_embedding_function(model_name: str):
        if model_name in {"", "local-hash", "hash", "offline"}:
            return LocalHashEmbeddingFunction()

        if model_name in {"chromadb-default", "default"}:
            logger.warning(
                "ChromaDB default embedding may download ONNX model files. "
                "Use 'local-hash' for offline startup."
            )
            return embedding_functions.DefaultEmbeddingFunction()

        return embedding_functions.SentenceTransformerEmbeddingFunction(
            model_name=model_name
        )

    @staticmethod
    def _default_collection_name(embedding_model_name: str) -> str:
        safe_model = re.sub(r"[^A-Za-z0-9_]+", "_", embedding_model_name or "local_hash")
        return f"gaokao_master_{safe_model[:40]}"
